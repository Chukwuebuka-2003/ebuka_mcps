"""
Test script for file processing capabilities.
Demonstrates PDF and DOCX text extraction and storage in RAG system.
"""

from utils.file_processor import FileProcessor
from rag.system import TutoringRAGSystem
import base64
import json


def test_pdf_extraction():
    """Test PDF text extraction."""
    print("\n" + "=" * 50)
    print("Testing PDF Extraction")
    print("=" * 50)

    # Initialize services
    rag_system = TutoringRAGSystem()
    processor = FileProcessor(rag_system)

    # Read a sample PDF file
    try:
        with open("sample.pdf", "rb") as f:
            pdf_content = f.read()

        print(f"✓ Loaded PDF file ({len(pdf_content)} bytes)")

        # Extract text preview
        preview = processor.extract_text_preview(
            pdf_content, "sample.pdf", max_chars=300
        )

        if preview["status"] == "success":
            print(f"✓ Extraction successful")
            print(f"  Total characters: {preview['total_characters']}")
            print(f"  Total pages: {preview['metadata'].get('total_pages', 'N/A')}")
            print(f"\n  Preview:\n  {preview['preview'][:200]}...")
        else:
            print(f"✗ Extraction failed: {preview.get('message')}")

    except FileNotFoundError:
        print("✗ sample.pdf not found. Please add a PDF file to test.")
    except Exception as e:
        print(f"✗ Error: {str(e)}")


def test_docx_extraction():
    """Test DOCX text extraction."""
    print("\n" + "=" * 50)
    print("Testing DOCX Extraction")
    print("=" * 50)

    # Initialize services
    rag_system = TutoringRAGSystem()
    processor = FileProcessor(rag_system)

    # Read a sample DOCX file
    try:
        with open("sample.docx", "rb") as f:
            docx_content = f.read()

        print(f"✓ Loaded DOCX file ({len(docx_content)} bytes)")

        # Extract text preview
        preview = processor.extract_text_preview(
            docx_content, "sample.docx", max_chars=300
        )

        if preview["status"] == "success":
            print(f"✓ Extraction successful")
            print(f"  Total characters: {preview['total_characters']}")
            print(
                f"  Total paragraphs: {preview['metadata'].get('total_paragraphs', 'N/A')}"
            )
            print(f"  Total tables: {preview['metadata'].get('total_tables', 'N/A')}")
            print(f"\n  Preview:\n  {preview['preview'][:200]}...")
        else:
            print(f"✗ Extraction failed: {preview.get('message')}")

    except FileNotFoundError:
        print("✗ sample.docx not found. Please add a DOCX file to test.")
    except Exception as e:
        print(f"✗ Error: {str(e)}")


def test_full_processing_and_storage():
    """Test complete file processing pipeline with RAG storage."""
    print("\n" + "=" * 50)
    print("Testing Full Processing Pipeline")
    print("=" * 50)

    # Initialize services
    rag_system = TutoringRAGSystem()
    processor = FileProcessor(rag_system)

    # Test with PDF
    try:
        with open("sample.pdf", "rb") as f:
            pdf_content = f.read()

        print("\nProcessing PDF...")
        result = processor.process_and_store_file(
            file_content=pdf_content,
            filename="sample.pdf",
            student_id="test_student_001",
            subject="Mathematics",
            topic="Calculus",
            difficulty_level=7,
            additional_metadata={"source": "textbook_chapter"},
        )

        if result["status"] == "success":
            print(f"✓ PDF processed and stored successfully")
            print(f"  Chunks stored: {result['chunks_stored']}")
            print(f"  Total characters: {result['total_characters']}")
            print(
                f"  Document IDs: {result['document_ids'][:2]}... ({len(result['document_ids'])} total)"
            )
        else:
            print(f"✗ Processing failed: {result.get('message')}")

    except FileNotFoundError:
        print("✗ sample.pdf not found")
    except Exception as e:
        print(f"✗ Error: {str(e)}")

    # Test with DOCX
    try:
        with open("sample.docx", "rb") as f:
            docx_content = f.read()

        print("\nProcessing DOCX...")
        result = processor.process_and_store_file(
            file_content=docx_content,
            filename="sample.docx",
            student_id="test_student_001",
            subject="History",
            topic="World War II",
            difficulty_level=6,
            additional_metadata={"source": "study_notes"},
        )

        if result["status"] == "success":
            print(f"✓ DOCX processed and stored successfully")
            print(f"  Chunks stored: {result['chunks_stored']}")
            print(f"  Total characters: {result['total_characters']}")
            print(
                f"  Document IDs: {result['document_ids'][:2]}... ({len(result['document_ids'])} total)"
            )
        else:
            print(f"✗ Processing failed: {result.get('message')}")

    except FileNotFoundError:
        print("✗ sample.docx not found")
    except Exception as e:
        print(f"✗ Error: {str(e)}")


def test_retrieval_after_processing():
    """Test retrieving stored file content from RAG system."""
    print("\n" + "=" * 50)
    print("Testing Retrieval After Processing")
    print("=" * 50)

    rag_system = TutoringRAGSystem()

    # Try to retrieve content we just stored
    print("\nQuerying for calculus content...")
    response = rag_system.generate_personalized_response(
        student_id="test_student_001",
        current_question="What did I learn about calculus?",
        subject="Mathematics",
        topic="Calculus",
        context_limit=3,
    )

    print(f"\n✓ Retrieved response:")
    print(f"  {response[:300]}...")


def test_base64_encoding():
    """Test base64 encoding for MCP tool usage."""
    print("\n" + "=" * 50)
    print("Testing Base64 Encoding (for MCP)")
    print("=" * 50)

    try:
        with open("sample.pdf", "rb") as f:
            pdf_content = f.read()

        # Encode to base64 (as required by MCP tool)
        encoded = base64.b64encode(pdf_content).decode("utf-8")

        print(f"✓ Original size: {len(pdf_content)} bytes")
        print(f"✓ Encoded size: {len(encoded)} characters")
        print(f"✓ First 100 chars: {encoded[:100]}...")

        # Decode back
        decoded = base64.b64decode(encoded)

        if decoded == pdf_content:
            print(f"✓ Encoding/decoding verified successfully")
        else:
            print(f"✗ Encoding/decoding mismatch!")

    except FileNotFoundError:
        print("✗ sample.pdf not found")
    except Exception as e:
        print(f"✗ Error: {str(e)}")


def create_sample_docx():
    """Create a sample DOCX file for testing."""
    print("\n" + "=" * 50)
    print("Creating Sample DOCX File")
    print("=" * 50)

    try:
        from docx import Document

        doc = Document()
        doc.add_heading("Calculus Study Notes", 0)

        doc.add_heading("Chapter 1: Limits", level=1)
        doc.add_paragraph(
            "A limit is the value that a function approaches as the input "
            "approaches some value. Limits are essential to calculus and are "
            "used to define continuity, derivatives, and integrals."
        )

        doc.add_heading("Key Concepts", level=2)
        doc.add_paragraph(
            "One-sided limits: approaching from left or right", style="List Bullet"
        )
        doc.add_paragraph(
            "Two-sided limits: both sides approach same value", style="List Bullet"
        )
        doc.add_paragraph(
            "Infinite limits: function grows without bound", style="List Bullet"
        )

        doc.add_heading("Example Problem", level=2)
        doc.add_paragraph(
            "Find the limit of f(x) = (x² - 1)/(x - 1) as x approaches 1.\n"
            "Solution: Factor the numerator to get (x+1)(x-1)/(x-1), "
            "cancel (x-1), and evaluate at x=1 to get 2."
        )

        doc.save("sample.docx")
        print("✓ Created sample.docx successfully")

    except ImportError:
        print("✗ python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        print(f"✗ Error: {str(e)}")


if __name__ == "__main__":
    print("\n" + "=" * 70)
    print(" FILE PROCESSOR TEST SUITE")
    print("=" * 70)

    # Create sample DOCX if it doesn't exist
    import os

    if not os.path.exists("sample.docx"):
        create_sample_docx()

    # Run tests
    test_pdf_extraction()
    test_docx_extraction()
    test_full_processing_and_storage()
    test_retrieval_after_processing()
    test_base64_encoding()

    print("\n" + "=" * 70)
    print(" TESTS COMPLETE")
    print("=" * 70)
    print("\nNote: Some tests may show 'file not found' if you don't have")
    print("sample PDF/DOCX files. Use create_sample_docx() to create a test file.")
