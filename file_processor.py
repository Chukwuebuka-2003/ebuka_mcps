"""
File processor for extracting text from PDFs and DOCX files.
Integrates with the RAG system to store extracted content.
"""

import io
import json
from typing import Dict, Any, Optional
from datetime import datetime
import pymupdf  # PyMuPDF
from docx import Document
from rag.types import LearningContext, MemoryType
from rag.system import TutoringRAGSystem


class FileProcessor:
    """
    Processes uploaded files (PDF, DOCX) and extracts text content
    for integration with the RAG system.
    """

    def __init__(self, rag_system: TutoringRAGSystem):
        """
        Initialize the file processor.

        Args:
            rag_system: Instance of TutoringRAGSystem for storing extracted content
        """
        self.rag_system = rag_system

    def extract_text_from_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """
        Extract text from a PDF file using PyMuPDF.

        Args:
            file_content: PDF file content as bytes

        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            # Open PDF from bytes
            pdf_document = pymupdf.open(stream=file_content, filetype="pdf")

            extracted_text = []
            metadata = {
                "total_pages": pdf_document.page_count,
                "page_texts": [],
            }

            # Extract text from each page
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                page_text = page.get_text()

                if page_text.strip():  # Only include non-empty pages
                    extracted_text.append(f"[Page {page_num + 1}]\n{page_text}")
                    metadata["page_texts"].append(
                        {
                            "page_number": page_num + 1,
                            "text_length": len(page_text),
                            "has_images": len(page.get_images()) > 0,
                        }
                    )

            # Close the document
            pdf_document.close()

            full_text = "\n\n".join(extracted_text)

            return {
                "status": "success",
                "text": full_text,
                "metadata": metadata,
                "total_characters": len(full_text),
            }

        except Exception as e:
            return {
                "status": "error",
                "message": f"Failed to extract text from PDF: {str(e)}",
            }

    def extract_text_from_docx(self, file_content: bytes) -> Dict[str, Any]:
        """
        Extract text from a DOCX file.

        Args:
            file_content: DOCX file content as bytes

        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            # Open DOCX from bytes
            docx_stream = io.BytesIO(file_content)
            document = Document(docx_stream)

            extracted_text = []
            metadata = {
                "total_paragraphs": len(document.paragraphs),
                "total_tables": len(document.tables),
            }

            # Extract text from paragraphs
            for para in document.paragraphs:
                if para.text.strip():
                    extracted_text.append(para.text)

            # Extract text from tables
            table_texts = []
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        table_texts.append(row_text)

            if table_texts:
                extracted_text.append("\n[Tables]\n" + "\n".join(table_texts))

            full_text = "\n\n".join(extracted_text)

            return {
                "status": "success",
                "text": full_text,
                "metadata": metadata,
                "total_characters": len(full_text),
            }

        except Exception as e:
            return {
                "status": "error",
                "message": f"Failed to extract text from DOCX: {str(e)}",
            }

    def process_and_store_file(
        self,
        file_content: bytes,
        filename: str,
        student_id: str,
        subject: str,
        topic: Optional[str] = None,
        difficulty_level: int = 5,
        additional_metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Process a file (PDF or DOCX) and store extracted content in the RAG system.

        Args:
            file_content: File content as bytes
            filename: Original filename
            student_id: Student identifier
            subject: Subject area
            topic: Optional specific topic
            difficulty_level: Difficulty level (1-10)
            additional_metadata: Optional additional metadata

        Returns:
            Dictionary with processing results and stored document IDs
        """
        try:
            # Determine file type
            file_extension = filename.lower().split(".")[-1]

            # Extract text based on file type
            if file_extension == "pdf":
                extraction_result = self.extract_text_from_pdf(file_content)
            elif file_extension in ["docx", "doc"]:
                extraction_result = self.extract_text_from_docx(file_content)
            else:
                return {
                    "status": "error",
                    "message": f"Unsupported file type: {file_extension}",
                }

            if extraction_result["status"] == "error":
                return extraction_result

            extracted_text = extraction_result["text"]
            file_metadata = extraction_result["metadata"]

            # Split text into chunks if it's too long (optional but recommended)
            chunks = self._chunk_text(extracted_text, chunk_size=1000, overlap=200)

            # Store each chunk in the RAG system
            stored_doc_ids = []

            for i, chunk in enumerate(chunks):
                # Prepare metadata - Pinecone only accepts simple types
                # Convert complex metadata to JSON strings
                metadata = {
                    "filename": filename,
                    "file_type": file_extension,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "extraction_metadata": json.dumps(
                        file_metadata
                    ),  # Serialize to JSON string
                    **(additional_metadata or {}),
                }

                # Ensure all metadata values are simple types
                for key, value in list(metadata.items()):
                    if isinstance(value, (dict, list)) and not isinstance(value, str):
                        metadata[key] = json.dumps(value)

                # Create learning context
                context = LearningContext(
                    student_id=student_id,
                    subject=subject,
                    topic=topic or f"uploaded_content_{filename}",
                    difficulty_level=difficulty_level,
                    learning_style="document_upload",
                    timestamp=datetime.now(),
                    content=chunk,
                    memory_type=MemoryType.CONTENT_MASTERY,
                    metadata=metadata,
                )

                # Store in RAG system
                doc_id = self.rag_system.store_learning_interaction(context)
                stored_doc_ids.append(doc_id)

            return {
                "status": "success",
                "message": f"Successfully processed and stored {len(chunks)} chunks from {filename}",
                "filename": filename,
                "file_type": file_extension,
                "total_characters": extraction_result["total_characters"],
                "chunks_stored": len(chunks),
                "document_ids": stored_doc_ids,
                "extraction_metadata": file_metadata,
            }

        except Exception as e:
            return {
                "status": "error",
                "message": f"Failed to process file: {str(e)}",
            }

    def _chunk_text(
        self, text: str, chunk_size: int = 1000, overlap: int = 200
    ) -> list[str]:
        """
        Split text into overlapping chunks for better retrieval.

        Args:
            text: Text to chunk
            chunk_size: Maximum characters per chunk
            overlap: Number of characters to overlap between chunks

        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            # Get chunk
            end = start + chunk_size
            chunk = text[start:end]

            # Try to break at sentence or paragraph boundary
            if end < len(text):
                # Look for sentence endings
                last_period = chunk.rfind(". ")
                last_newline = chunk.rfind("\n")

                break_point = max(last_period, last_newline)

                if break_point > chunk_size * 0.5:  # Only break if we're past halfway
                    chunk = text[start : start + break_point + 1]
                    end = start + break_point + 1

            chunks.append(chunk.strip())

            # Move start position with overlap
            start = end - overlap

        return chunks

    def extract_text_preview(
        self, file_content: bytes, filename: str, max_chars: int = 500
    ) -> Dict[str, Any]:
        """
        Extract a preview of text from a file without storing it.
        Useful for showing users what was extracted before storing.

        Args:
            file_content: File content as bytes
            filename: Original filename
            max_chars: Maximum characters to return in preview

        Returns:
            Dictionary with preview text and metadata
        """
        try:
            file_extension = filename.lower().split(".")[-1]

            if file_extension == "pdf":
                extraction_result = self.extract_text_from_pdf(file_content)
            elif file_extension in ["docx", "doc"]:
                extraction_result = self.extract_text_from_docx(file_content)
            else:
                return {
                    "status": "error",
                    "message": f"Unsupported file type: {file_extension}",
                }

            if extraction_result["status"] == "error":
                return extraction_result

            full_text = extraction_result["text"]
            preview = full_text[:max_chars]

            if len(full_text) > max_chars:
                preview += "..."

            return {
                "status": "success",
                "preview": preview,
                "total_characters": len(full_text),
                "metadata": extraction_result["metadata"],
            }

        except Exception as e:
            return {
                "status": "error",
                "message": f"Failed to extract preview: {str(e)}",
            }
