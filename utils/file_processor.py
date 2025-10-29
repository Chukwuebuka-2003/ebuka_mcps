import io
import json
import logging
from datetime import datetime
from typing import Any, Dict, List, Optional
import os

import pymupdf  # type: ignore
from docx import Document
from rag.system import TutoringRAGSystem
from rag.types import LearningContext, MemoryType
from openai import OpenAI

logger = logging.getLogger(__name__)

# Initialize OpenAI client for AI subject detection
openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


class FileProcessor:
    """
    Processes uploaded files (PDF, DOCX) and extracts text content
    for integration with the RAG system.
    """

    def __init__(self, rag_system: TutoringRAGSystem):
        """
        Initialize the file processor.

        Args:
            rag_system: An instance of TutoringRAGSystem for storing learning interactions.
        """
        self.rag_system = rag_system

    @staticmethod
    def detect_subject_from_content(text: str) -> str:
        """
        Use AI to intelligently detect the subject from document content.

        Args:
            text: The document text content to analyze

        Returns:
            Detected subject name (e.g., "Mathematics", "Physics", "Computer Science")
        """
        try:
            # Use first 2000 characters for analysis
            sample_text = text[:2000]

            prompt = f"""Analyze the following document content and identify the primary academic subject.

Content: {sample_text}

Return ONLY the subject name from this list:
- Mathematics
- Physics
- Chemistry
- Biology
- Computer Science
- English
- History
- Geography
- Economics
- Psychology
- Philosophy
- Art
- Music
- Engineering
- Business
- General

If the content clearly fits multiple subjects, choose the most dominant one.
If unclear, return "General".

Subject:"""

            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are an expert at identifying academic subjects from document content. Always respond with a single subject name."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=20
            )

            detected_subject = response.choices[0].message.content.strip()
            logger.info(f"🤖 AI detected subject from document: {detected_subject}")
            return detected_subject

        except Exception as e:
            logger.error(f"AI subject detection from document failed: {e}")
            return "General"

    def extract_text_from_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """
        Extract text from a PDF file.

        Args:
            file_content: PDF file content as bytes

        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            # Open PDF from bytes
            pdf_document = pymupdf.open(stream=file_content, filetype="pdf")

            extracted_text = []
            metadata = {
                "total_pages": pdf_document.page_count,
                "page_texts": [],
            }

            # Extract text from each page
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                page_text = page.get_text()

                if page_text.strip():  # Only include non-empty pages
                    extracted_text.append(f"[Page {page_num + 1}]\n{page_text}")
                    metadata["page_texts"].append(
                        {
                            "page_number": page_num + 1,
                            "text_length": len(page_text),
                            "has_images": len(page.get_images()) > 0,
                        }
                    )

            # Close the document
            pdf_document.close()

            full_text = "\n\n".join(extracted_text)

            return {
                "status": "success",
                "text": full_text,
                "metadata": metadata,
                "total_characters": len(full_text),
            }

        except Exception as e:
            return {
                "status": "error",
                "message": f"Failed to extract text from PDF: {str(e)}",
            }

    def extract_text_from_docx(self, file_content: bytes) -> Dict[str, Any]:
        """
        Extract text from a DOCX file.

        Args:
            file_content: DOCX file content as bytes

        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            # Open DOCX from bytes
            docx_stream = io.BytesIO(file_content)
            document = Document(docx_stream)

            extracted_text = []
            metadata = {
                "total_paragraphs": len(document.paragraphs),
                "total_tables": len(document.tables),
            }

            # Extract text from paragraphs
            for para in document.paragraphs:
                if para.text.strip():
                    extracted_text.append(para.text)

            # Extract text from tables
            table_texts = []
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        table_texts.append(row_text)

            if table_texts:
                extracted_text.append("\n[Tables]\n" + "\n".join(table_texts))

            full_text = "\n\n".join(extracted_text)

            return {
                "status": "success",
                "text": full_text,
                "metadata": metadata,
                "total_characters": len(full_text),
            }

        except Exception as e:
            return {
                "status": "error",
                "message": f"Failed to extract text from DOCX: {str(e)}",
            }

    def _chunk_text(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """
        Splits text into chunks with a specified overlap.
        """
        if not text:
            return []

        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            if end >= len(text):
                break
            start += chunk_size - overlap
        return chunks

    def process_and_store_file(
        self,
        file_content: bytes,
        filename: str,
        student_id: str,
        subject: str,
        topic: Optional[str] = None,
        difficulty_level: int = 5,
        additional_metadata: Optional[Dict[str, Any]] = None,
        document_title: Optional[str] = None,  # Allow custom title
    ) -> Dict[str, Any]:
        """
        Process a file (PDF or DOCX) and store extracted content in the RAG system.

        Args:
            file_content: File content as bytes
            filename: Original filename
            student_id: Student identifier
            subject: Subject area
            topic: Optional specific topic
            difficulty_level: Difficulty level (1-10)
            additional_metadata: Optional additional metadata
            document_title: Optional custom title for citations

        Returns:
            Dictionary with processing results and stored document IDs
        """
        try:
            # Determine file type
            file_extension = filename.lower().split(".")[-1]

            # Extract text based on file type
            if file_extension == "pdf":
                extraction_result = self.extract_text_from_pdf(file_content)
            elif file_extension in ["docx", "doc"]:
                extraction_result = self.extract_text_from_docx(file_content)
            else:
                return {
                    "status": "error",
                    "message": f"Unsupported file type: {file_extension}",
                }

            if extraction_result["status"] == "error":
                return extraction_result

            extracted_text = extraction_result["text"]
            file_metadata = extraction_result["metadata"]

            # Use AI to detect subject if "General" was provided or if we want to override
            if subject == "General" or subject is None:
                detected_subject = self.detect_subject_from_content(extracted_text)
                logger.info(f"📚 AI detected subject '{detected_subject}' from document content")
                subject = detected_subject
            else:
                logger.info(f"📚 Using provided subject: {subject}")

            # Split text into chunks if it's too long (optional but recommended)
            chunks = self._chunk_text(extracted_text, chunk_size=1000, overlap=200)

            # Use custom title or filename as fallback
            doc_title = document_title or filename

            # Store each chunk in the RAG system
            stored_doc_ids = []

            for i, chunk in enumerate(chunks):
                # Prepare metadata - Pinecone only accepts simple types
                # Convert complex metadata to JSON strings
                metadata = {
                    "filename": filename,
                    "document_title": doc_title,  # Add document title for citations
                    "file_type": file_extension,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "extraction_metadata": json.dumps(
                        file_metadata
                    ),  # Serialize to JSON string
                    **(additional_metadata or {}),
                }

                # Ensure all metadata values are simple types
                for key, value in list(metadata.items()):
                    if isinstance(value, (dict, list)) and not isinstance(value, str):
                        metadata[key] = json.dumps(value)

                # Create learning context
                context = LearningContext(
                    student_id=student_id,
                    subject=subject,
                    topic=topic or f"uploaded_content_{filename}",
                    difficulty_level=difficulty_level,
                    learning_style="document_upload",
                    timestamp=datetime.now(),
                    content=chunk,
                    memory_type=MemoryType.CONTENT_MASTERY,
                    metadata=metadata,
                    document_title=doc_title,
                )

                # Store in RAG system
                doc_id = self.rag_system.store_learning_interaction(context)
                stored_doc_ids.append(doc_id)

            return {
                "status": "success",
                "message": f"Successfully processed and stored {len(stored_doc_ids)} document chunks.",
                "document_ids": stored_doc_ids,
                "total_characters": extraction_result["total_characters"],
                "chunks_stored": len(stored_doc_ids),
                "detected_subject": subject,  # Return the detected/used subject
                "document_title": doc_title,
            }

        except Exception as e:
            logger.exception(f"Error processing and storing file {filename}: {e}")
            return {"status": "error", "message": f"File processing failed: {str(e)}"}
